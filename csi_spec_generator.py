"""
CSI Specification Generator for Section 23 51 10
MECHANICAL DRAFT SYSTEMS FOR COMMERCIAL/RESIDENTIAL VENTING
Compliant with UL 705, UL 378, NFPA 54, NFPA 211, IMC, IFGC
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

class CSISpecificationGenerator:
    """Generates CSI MasterFormat Section 23 51 10 specifications"""
    
    def __init__(self):
        self.section_number = "23 51 10"
        self.section_title = "MECHANICAL DRAFT SYSTEMS"
        
    def generate_specification(self, project_info, products_selected, system_data):
        """Generate complete CSI specification document"""
        doc = Document()
        self._setup_document_styles(doc)
        self._add_header(doc, project_info)
        self._add_part_1_general(doc)
        self._add_part_2_products(doc, products_selected, system_data)
        self._add_part_3_execution(doc, products_selected)
        return doc
    
    def _setup_document_styles(self, doc):
        """Set up document styles"""
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
    
    def _add_header(self, doc, project_info):
        """Add specification header"""
        p = doc.add_paragraph()
        p.add_run(f"SECTION {self.section_number}").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph()
        p.add_run(self.section_title).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
        doc.add_paragraph(f"Project: {project_info.get('name', 'TBD')}")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)
        doc.add_paragraph()
    
    def _add_section_heading(self, doc, number, title):
        """Add section heading"""
        p = doc.add_paragraph()
        run = p.add_run(f"{number} {title}")
        run.bold = True
        run.underline = True
    
    def _add_part_1_general(self, doc):
        """Add PART 1 - GENERAL"""
        p = doc.add_paragraph("PART 1 - GENERAL")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        doc.add_paragraph()
        
        # 1.1 SUMMARY
        self._add_section_heading(doc, "1.1", "SUMMARY")
        doc.add_paragraph("A. Furnish and install powered draft equipment, electronic controls, and accessories for commercial/residential venting systems.")
        doc.add_paragraph()
        
        # 1.2 REFERENCES
        self._add_section_heading(doc, "1.2", "REFERENCES")
        doc.add_paragraph("A. UL 705 - Power Ventilators")
        doc.add_paragraph("B. UL 378 - Draft Equipment")
        doc.add_paragraph("C. NFPA 54 - National Fuel Gas Code")
        doc.add_paragraph("D. NFPA 211 - Chimneys, Fireplaces, Vents")
        doc.add_paragraph("E. IMC - International Mechanical Code")
        doc.add_paragraph("F. IFGC - International Fuel Gas Code")
        doc.add_paragraph()
        
        # 1.3 SUBMITTALS
        self._add_section_heading(doc, "1.3", "SUBMITTALS")
        doc.add_paragraph("A. Product data: Manufacturer's specifications, performance curves, and electrical data.")
        doc.add_paragraph("B. Shop drawings: Equipment layout, wiring diagrams, and sequence of operations.")
        doc.add_paragraph("C. O&M manuals: Installation instructions, maintenance procedures, and parts lists.")
        doc.add_paragraph("D. Startup report: Factory commissioning report with test data.")
        doc.add_paragraph()
        
        # 1.4 QUALITY ASSURANCE
        self._add_section_heading(doc, "1.4", "QUALITY ASSURANCE")
        doc.add_paragraph("A. Equipment shall be UL 705 and UL 378 Listed.")
        doc.add_paragraph("B. Installer shall be certified by equipment manufacturer.")
        doc.add_paragraph("C. Factory startup service required by manufacturer's representative.")
        doc.add_paragraph()
        
        # 1.5 WARRANTY
        self._add_section_heading(doc, "1.5", "WARRANTY")
        doc.add_paragraph("A. Two (2) year manufacturer's warranty covering defects in materials and workmanship.")
        doc.add_paragraph()
    
    def _add_part_2_products(self, doc, products, system_data):
        """Add PART 2 - PRODUCTS"""
        p = doc.add_paragraph("PART 2 - PRODUCTS")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        doc.add_paragraph()
        
        # 2.1 MANUFACTURER
        self._add_section_heading(doc, "2.1", "MANUFACTURER")
        doc.add_paragraph("A. US Draft Co., a division of R.M. Manifold Group, Inc.")
        doc.add_paragraph("   100 S Sylvania Ave, Fort Worth, TX 76111")
        doc.add_paragraph("   Phone: 817-393-4029 | www.usdraft.com")
        doc.add_paragraph()
        
        # 2.2 DRAFT INDUCER
        if products.get('draft_inducer'):
            self._add_section_heading(doc, "2.2", "POWERED DRAFT INDUCER")
            self._add_draft_inducer_spec(doc, products['draft_inducer'], system_data)
        
        # 2.3 CONTROLLER
        if products.get('controller'):
            section = "2.3" if products.get('draft_inducer') else "2.2"
            self._add_section_heading(doc, section, "ELECTRONIC CONTROL SYSTEM")
            self._add_controller_spec(doc, products['controller'])
        
        # 2.4 CDS3 CHIMNEY DRAFT STABILIZATION
        if products.get('cds3'):
            section_num = "2.2"
            if products.get('draft_inducer'):
                section_num = "2.3"
            if products.get('controller'):
                section_num = "2.4"
            
            self._add_section_heading(doc, section_num, "CDS3 CHIMNEY DRAFT STABILIZATION SYSTEM")
            self._add_cds3_spec(doc, products, system_data)
        
        # 2.5 ODCS (for non-Category IV systems)
        if products.get('odcs'):
            section_num = "2.2"
            if products.get('draft_inducer'):
                section_num = "2.3"
            if products.get('controller'):
                section_num = "2.4"
            if products.get('cds3'):
                section_num = "2.5"
            
            self._add_section_heading(doc, section_num, "ODCS OVERDRAFT CONTROL SYSTEM")
            doc.add_paragraph("A. ODCS with motorized damper, pressure sensor, and controller interface.")
            doc.add_paragraph("B. UL 378 Listed for automatic control of excessive draft on natural draft appliances.")
            doc.add_paragraph("C. Requires electronic control system (V150/V250/V350) for operation.")
            doc.add_paragraph()
    
    def _add_draft_inducer_spec(self, doc, inducer, system_data):
        """Add draft inducer specification"""
        series = inducer.get('series', '')
        model = inducer.get('model', '')
        
        # Get CFM and pressure, format properly
        cfm = system_data.get('cfm', 'TBD')
        if isinstance(cfm, (int, float)):
            cfm_str = f"{int(round(cfm))}"
        else:
            cfm_str = str(cfm)
        
        pressure = system_data.get('static_pressure', 'TBD')
        if isinstance(pressure, (int, float)):
            pressure_str = f"{pressure:.2f}"
        else:
            pressure_str = str(pressure)
        
        # Series names
        series_names = {
            'TRV': 'TRV Series True Inline Draft Inducer',
            'T9F': 'T9F Series 90-Degree Inline Draft Inducer',
            'CBX': 'CBX Series Termination Mount Chimney Fan'
        }
        
        series_name = series_names.get(series, 'Draft Inducer')
        
        doc.add_paragraph(f"A. {series_name}, Model {model}")
        doc.add_paragraph(f"B. UL 705 Listed Powered Draft Equipment")
        doc.add_paragraph(f"C. Design Airflow: {cfm_str} CFM at {pressure_str} inches w.c. static pressure")
        
        # Material selection
        cat = system_data.get('appliance_category', 'I')
        if cat in ['II', 'IV']:
            material = "316L stainless steel for condensing flue gas applications"
        else:
            material = "Aluminized steel or 316L stainless steel"
        
        doc.add_paragraph(f"D. Construction: {material} housing, aluminum impeller, EC motor")
        doc.add_paragraph("E. Motor: Electronically commutated (EC) permanent magnet, IE4 efficiency, IP54 rated")
        doc.add_paragraph("F. Temperature Rating: 550°F continuous, 650°F intermittent")
        doc.add_paragraph("G. Control: 0-10VDC modulating input, EC-Flow™ constant airflow technology")
        doc.add_paragraph("H. Features: Vibration isolation, integral flow proving switch, UL 705 overload protection")
        doc.add_paragraph()
    
    def _add_cds3_spec(self, doc, products, system_data):
        """Add CDS3 specification with datasheet details"""
        num_appliances = len(system_data.get('appliances', []))
        
        doc.add_paragraph("A. CDS3 Self-Contained Chimney Draft Stabilization System")
        doc.add_paragraph(f"B. Quantity: {num_appliances} unit(s) - one per appliance connector")
        doc.add_paragraph("C. Application: Category IV condensing appliances with low draft requirements")
        doc.add_paragraph()
        
        doc.add_paragraph("D. System Components (each CDS3 unit includes):")
        doc.add_paragraph("   1. Motorized Damper:")
        doc.add_paragraph("      a. 24VAC actuator with 2-second stroke time")
        doc.add_paragraph("      b. Spring return to fail-safe position")
        doc.add_paragraph("      c. Bi-directional modulation (0-100%)")
        doc.add_paragraph("      d. Stainless steel construction for condensing applications")
        doc.add_paragraph()
        
        doc.add_paragraph("   2. Pressure Transducer:")
        doc.add_paragraph("      a. Bidirectional measurement range: ±2.0 in w.c.")
        doc.add_paragraph("      b. Resolution: 0.001 in w.c.")
        doc.add_paragraph("      c. Accuracy: ±1% of reading")
        doc.add_paragraph("      d. Temperature compensated")
        doc.add_paragraph()
        
        doc.add_paragraph("   3. Integrated PID Controller:")
        doc.add_paragraph("      a. Self-contained microprocessor control")
        doc.add_paragraph("      b. Auto-tuning PID algorithm")
        doc.add_paragraph("      c. Field-adjustable setpoint: -0.10 to -0.01 in w.c.")
        doc.add_paragraph("      d. No external controller required")
        doc.add_paragraph()
        
        doc.add_paragraph("E. Performance:")
        doc.add_paragraph("   1. Maintains constant draft pressure regardless of variations")
        doc.add_paragraph("   2. Response time: Less than 2 seconds to pressure changes")
        doc.add_paragraph("   3. Prevents excessive draft that wastes energy")
        doc.add_paragraph("   4. Maintains stable combustion conditions")
        doc.add_paragraph()
        
        doc.add_paragraph("F. Electrical:")
        doc.add_paragraph("   1. Power: 24VAC transformer (included)")
        doc.add_paragraph("   2. Power consumption: 15VA maximum")
        doc.add_paragraph()
        
        doc.add_paragraph("G. Installation:")
        doc.add_paragraph("   1. Mounts in appliance connector (breeching)")
        doc.add_paragraph("   2. Install between appliance outlet and common vent")
        doc.add_paragraph("   3. Pressure tap at appliance flue collar")
        doc.add_paragraph("   4. Operates independently - no inter-unit communication required")
        doc.add_paragraph()
        
        doc.add_paragraph("H. Compliance:")
        doc.add_paragraph("   1. UL 378 Listed - Draft Equipment")
        doc.add_paragraph("   2. Complies with NFPA 54, NFPA 211, IMC, IFGC")
        doc.add_paragraph()
    
    def _add_controller_spec(self, doc, controller):
        """Add controller specification"""
        model = controller.get('base_model', 'V250')
        full_model = controller.get('model', model)
        
        # Model specs
        model_specs = {
            'V150': ('LCD display', '1-2 appliances'),
            'V250': ('4-inch color touchscreen', '1-6 appliances'),
            'V350': ('7-inch color touchscreen', '1-15 appliances')
        }
        
        display, capacity = model_specs.get(model, ('Touchscreen', '1-6 appliances'))
        
        doc.add_paragraph(f"A. Model {full_model} Electronic Vent Control System")
        doc.add_paragraph(f"B. UL 378 Listed Draft Control Equipment")
        doc.add_paragraph(f"C. Display: {display}")
        doc.add_paragraph(f"D. Capacity: Controls {capacity}")
        doc.add_paragraph("E. Safety Interlocking per NFPA 54 Section 13.2:")
        doc.add_paragraph("   1. Appliance proving via differential pressure switch")
        doc.add_paragraph("   2. Draft proving before appliance enable")
        doc.add_paragraph("   3. Flow proving from draft inducer")
        doc.add_paragraph("   4. Adjustable pre-purge and post-purge timers")
        doc.add_paragraph("   5. Low draft cutout protection")
        doc.add_paragraph("F. Features: EC-Flow™ control, data logging, alarm outputs, password protection")
        
        if model in ['V250', 'V350']:
            doc.add_paragraph("G. Communication: Modbus RTU or BACnet MS/TP (optional)")
        
        doc.add_paragraph()
    
    def _add_part_3_execution(self, doc, products):
        """Add PART 3 - EXECUTION"""
        p = doc.add_paragraph("PART 3 - EXECUTION")
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(11)
        doc.add_paragraph()
        
        # 3.1 INSTALLATION
        self._add_section_heading(doc, "3.1", "INSTALLATION")
        doc.add_paragraph("A. Install per manufacturer's instructions, approved drawings, and applicable codes.")
        doc.add_paragraph("B. Provide structural support independent of vent piping.")
        doc.add_paragraph("C. Maintain manufacturer's recommended clearances for service access.")
        
        if products.get('draft_inducer'):
            series = products['draft_inducer'].get('series', '')
            if series == 'CBX':
                doc.add_paragraph("D. CBX Fan: Mount at top of stack with weatherproof flashing and sealing.")
            elif series == 'TRV':
                doc.add_paragraph("D. TRV Fan: Install inline with 3D upstream and 5D downstream straight duct.")
            elif series == 'T9F':
                doc.add_paragraph("D. T9F Fan: Install at 90° turn with vibration isolation and independent support.")
        
        doc.add_paragraph()
        
        # 3.2 ELECTRICAL
        self._add_section_heading(doc, "3.2", "ELECTRICAL INSTALLATION")
        doc.add_paragraph("A. Provide dedicated 120VAC/1Ph/60Hz, 15A circuit per NEC.")
        doc.add_paragraph("B. Install surge protection per IEEE C62.41.")
        doc.add_paragraph("C. Route control wiring per NEC Article 725 for Class 2 circuits.")
        doc.add_paragraph("D. Provide proper grounding per NEC Article 250.")
        doc.add_paragraph("E. Install disconnect switch within sight of equipment per NEC Article 430.102.")
        doc.add_paragraph()
        
        # 3.3 STARTUP
        self._add_section_heading(doc, "3.3", "STARTUP AND COMMISSIONING")
        doc.add_paragraph("A. Factory-trained technician shall perform startup and commissioning.")
        doc.add_paragraph("B. Verify all safety interlocks operate correctly.")
        doc.add_paragraph("C. Measure and record draft pressure and airflow.")
        doc.add_paragraph("D. Test system under all operating scenarios.")
        doc.add_paragraph("E. Submit commissioning report with test data to Engineer.")
        doc.add_paragraph()
        
        # 3.4 TRAINING
        self._add_section_heading(doc, "3.4", "TRAINING")
        doc.add_paragraph("A. Provide 4 hours of on-site training by factory representative.")
        doc.add_paragraph("B. Training topics: Operation, troubleshooting, maintenance, and safety procedures.")
        doc.add_paragraph()
        doc.add_paragraph()
        
        p = doc.add_paragraph("END OF SECTION 23 51 10")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True

